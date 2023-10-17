from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Name the colors
BRIGHT_GREEN = WD_COLOR_INDEX.BRIGHT_GREEN
YELLOW = WD_COLOR_INDEX.YELLOW
RED = WD_COLOR_INDEX.RED

# This is a function that generates a colored word document based on a list of strings [los] and
# a list of floats [lof], where [los] and [lof] are of equal length
def generate_docx(los: list[str], lof: list[float]):


    # Create a document
    document = Document()   
    # Add a paragraph to the document
    paragraph = document.add_paragraph()
    # Get the length of [los] or [lof]
    num = len(los)


    # When there is only one string
    if num == 1:

        # Capitalize the word
        los[0] = los[0].capitalize()

        # Add the highlighted word to the paragraph based on the float
        if 0.9 < lof[0] and lof[0] <= 1:
            run = paragraph.add_run(los[0]).font.highlight_color = BRIGHT_GREEN
        elif 0.8 < lof[0] and lof[0] <= 0.9:
            run = paragraph.add_run(los[0]).font.highlight_color = YELLOW
        elif -1 < lof[0] and lof[0] <= 0.8:
            run = paragraph.add_run(los[0]).font.highlight_color = RED
        # Set to no color if float out of range
        else:
            run = paragraph.add_run(los[0])


    # When there is more than one string
    else:

        # Loop from the first string to the second last string
        for i in range(0, num - 1, 1):

            if (los[i] != '(' and los[i] != '[' and los[i] != '{' and los[i] != '<' and los[i] != '``' and
            los[i + 1] != '.' and los[i + 1] != '?' and los[i + 1] != '!' and 
            los[i + 1] != ',' and los[i + 1] != ';' and los[i + 1] != ':' and
            los[i + 1] != ')' and los[i + 1] != ']' and los[i + 1] != '}' and los[i + 1] != '>' and los[i + 1] != '"' and
            los[i + 1][0] != "'"):

                # Capitalize if it is the first word
                if i == 0:
                    los[0] = los[0].capitalize()

                # Add the highlighted word to the paragraph followed by a space
                if 0.9 < lof[i] and lof[i] <= 1:
                    run = paragraph.add_run(los[i] + ' ').font.highlight_color = BRIGHT_GREEN
                elif 0.8 < lof[i] and lof[i] <= 0.9:
                    run = paragraph.add_run(los[i] + ' ').font.highlight_color = YELLOW
                elif -1 < lof[i] and lof[i] <= 0.8:
                    run = paragraph.add_run(los[i] + ' ').font.highlight_color = RED
                # Set to no color if float out of range
                else:
                    run = paragraph.add_run(los[i] + ' ')

            elif (los[i] == '(' or los[i] == '[' or los[i] == '{' or los[i] == '<' or los[i] == '``' or
            los[i + 1] == '.' or los[i + 1] == '?' or los[i + 1] == '!' or 
            los[i + 1] == ',' or los[i + 1] == ';' or los[i + 1] == ':' or
            los[i + 1] == ')' or los[i + 1] == ']' or los[i + 1] == '}' or los[i + 1] == '>' or los[i + 1] == '"' or
            los[i + 1][0] == "'"):

                # Capitalize if it is the first word
                if i == 0:
                    los[0] = los[0].capitalize()

                # Add the highlighted word to the paragraph
                if 0.9 < lof[i] and lof[i] <= 1:
                    run = paragraph.add_run(los[i]).font.highlight_color = BRIGHT_GREEN
                elif 0.8 < lof[i] and lof[i] <= 0.9:
                    run = paragraph.add_run(los[i]).font.highlight_color = YELLOW
                elif -1 < lof[i] and lof[i] <= 0.8:
                    run = paragraph.add_run(los[i]).font.highlight_color = RED
                # Set to no color if float out of range
                else:
                    run = paragraph.add_run(los[i])

        # Add the last word
        if 0.9 < lof[num - 1] and lof[num - 1] <= 1:
            run = paragraph.add_run(los[num - 1]).font.highlight_color = BRIGHT_GREEN
        elif 0.8 < lof[num - 1] and lof[num - 1] <= 0.9:
            run = paragraph.add_run(los[num - 1]).font.highlight_color = YELLOW
        elif -1 < lof[num - 1] and lof[num - 1] <= 0.8:
            run = paragraph.add_run(los[num - 1]).font.highlight_color = RED
        # Set to no color if float out of range
        else:
            run = paragraph.add_run(los[num - 1])
    

    # Save the colored document as "colored_docx.docx"
    document.save('colored_docx.docx')