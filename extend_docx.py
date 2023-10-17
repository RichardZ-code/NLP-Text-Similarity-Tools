from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Name the colors
BRIGHT_GREEN = WD_COLOR_INDEX.BRIGHT_GREEN
YELLOW = WD_COLOR_INDEX.YELLOW
RED = WD_COLOR_INDEX.RED

RED_VALUE = 0.3
YELLOW_VALUE = 0.7

"""
This is a function that adds a new paragraph to [doc] based on a list of strings [los] and 
a list of floats [lof], where [los] and [lof] are of equal length
"""
def partial_docx(los: list[str], lof: list[float], doc):

 
    # Add a paragraph to the document
    paragraph = doc.add_paragraph()
    # Get the length of [los] or [lof]
    num = len(los)


    # When there is only one string
    if num == 1:

        # Capitalize the word
        los[0] = los[0].capitalize()

        # Add the highlighted word to the paragraph based on the float
        if YELLOW_VALUE < lof[0] and lof[0] <= 1:
            run = paragraph.add_run(los[0]).font.highlight_color = BRIGHT_GREEN
        elif RED_VALUE < lof[0] and lof[0] <= YELLOW_VALUE:
            run = paragraph.add_run(los[0]).font.highlight_color = YELLOW
        elif -1 < lof[0] and lof[0] <= RED_VALUE:
            run = paragraph.add_run(los[0]).font.highlight_color = RED
        # Set to no color if float out of range
        else:
            run = paragraph.add_run(los[0])


    # When there is more than one string
    else:

        # Loop from the first string to the second last string
        for i in range(0, num - 1, 1):

            if (los[i] != '(' and los[i] != '[' and los[i] != '{' and los[i] != '<' and los[i] != '``' and
            los[i + 1] != '.' and los[i + 1] != '?' and los[i + 1] != '!' and 
            los[i + 1] != ',' and los[i + 1] != ';' and los[i + 1] != ':' and
            los[i + 1] != ')' and los[i + 1] != ']' and los[i + 1] != '}' and los[i + 1] != '>' and los[i + 1] != '"' and
            los[i + 1][0] != "'"):

                # Capitalize if it is the first word
                if i == 0:
                    los[0] = los[0].capitalize()

                # Add the highlighted word to the paragraph followed by a space
                if YELLOW_VALUE < lof[i] and lof[i] <= 1:
                    run = paragraph.add_run(los[i] + ' ').font.highlight_color = BRIGHT_GREEN
                elif RED_VALUE < lof[i] and lof[i] <= YELLOW_VALUE:
                    run = paragraph.add_run(los[i] + ' ').font.highlight_color = YELLOW
                elif -1 < lof[i] and lof[i] <= RED_VALUE:
                    run = paragraph.add_run(los[i] + ' ').font.highlight_color = RED
                # Set to no color if float out of range
                else:
                    run = paragraph.add_run(los[i] + ' ')

            elif (los[i] == '(' or los[i] == '[' or los[i] == '{' or los[i] == '<' or los[i] == '``' or
            los[i + 1] == '.' or los[i + 1] == '?' or los[i + 1] == '!' or 
            los[i + 1] == ',' or los[i + 1] == ';' or los[i + 1] == ':' or
            los[i + 1] == ')' or los[i + 1] == ']' or los[i + 1] == '}' or los[i + 1] == '>' or los[i + 1] == '"' or
            los[i + 1][0] == "'"):

                # Capitalize if it is the first word
                if i == 0:
                    los[0] = los[0].capitalize()

                # Add the highlighted word to the paragraph
                if YELLOW_VALUE < lof[i] and lof[i] <= 1:
                    run = paragraph.add_run(los[i]).font.highlight_color = BRIGHT_GREEN
                elif RED_VALUE < lof[i] and lof[i] <= YELLOW_VALUE:
                    run = paragraph.add_run(los[i]).font.highlight_color = YELLOW
                elif -1 < lof[i] and lof[i] <= RED_VALUE:
                    run = paragraph.add_run(los[i]).font.highlight_color = RED
                # Set to no color if float out of range
                else:
                    run = paragraph.add_run(los[i])

        # Add the last word
        if YELLOW_VALUE < lof[num - 1] and lof[num - 1] <= 1:
            run = paragraph.add_run(los[num - 1]).font.highlight_color = BRIGHT_GREEN
        elif RED_VALUE < lof[num - 1] and lof[num - 1] <= YELLOW_VALUE:
            run = paragraph.add_run(los[num - 1]).font.highlight_color = YELLOW
        elif -1 < lof[num - 1] and lof[num - 1] <= RED_VALUE:
            run = paragraph.add_run(los[num - 1]).font.highlight_color = RED
        # Set to no color if float out of range
        else:
            run = paragraph.add_run(los[num - 1])



"""
This is a function that generates a colored word document based on a list of list[str] [los] and 
a list of list[float] [lof], where each corresponding list[str] and list[float] are of equal length; 
[los] and [lof] are of equal length
"""
def whole_docx(los: list[list[str]], lof: list[list[float]], doc_name = 'color_doc.docx') -> None:


    # Create a document
    document = Document()


    # Count the number of list[str] in los
    count = len(los)
    # Loop from the first list[str] to the last list[str]
    for i in range(0, count, 1):
        partial_docx(los[i], lof[i], document)
    

    # Save the colored document as "doc_name"
    document.save(doc_name)